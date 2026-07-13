from collections.abc import AsyncIterator
from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4

import fitz
import pytest
from docx import Document as DocxDocument

from app.chats.models import Chat, Message, MessageRole, MessageStatus
from app.chats.schemas import ChatDocumentSelectionUpdate, MessageCreate
from app.chats.service import ChatService
from app.core.config import Settings
from app.core.exceptions import (
    DependencyUnavailableError,
    InvalidRequestError,
    UnsupportedMediaTypeError,
)
from app.documents.chunking import DocumentChunker
from app.documents.models import Document, DocumentStatus
from app.documents.parsers import DocxParser, ParsedBlock, ParsedDocument, PdfParser, TxtParser
from app.documents.service import DocumentService
from app.embeddings.base import EmbeddingTask


class FakeStorage:
    def __init__(self) -> None:
        self.uploaded: dict[str, bytes] = {}
        self.deleted: list[str] = []

    async def upload(self, key: str, data: bytes, content_type: str) -> None:
        del content_type
        self.uploaded[key] = data

    async def download(self, key: str) -> bytes:
        return self.uploaded[key]

    async def delete(self, key: str) -> None:
        self.deleted.append(key)

    async def exists(self, key: str) -> bool:
        return key in self.uploaded


class FakeJobQueue:
    def __init__(self, *, fail: bool = False) -> None:
        self.fail = fail
        self.enqueued: list[tuple[str, str]] = []

    async def enqueue_job(self, name: str, document_id: str) -> object:
        assert name == "process_document"
        assert document_id
        if self.fail:
            raise ConnectionError("Redis unavailable")
        self.enqueued.append((name, document_id))
        return object()


class FakeDocumentRepository:
    document: Document | None = None

    async def list_all(self, *, limit: int = 100, offset: int = 0) -> tuple[list[Document], int]:
        del limit, offset
        return [], 0

    async def get(self, document_id: object, *, with_chunks: bool = False) -> Document | None:
        del document_id, with_chunks
        return None

    async def create(
        self,
        *,
        original_filename: str,
        storage_key: str,
        mime_type: str,
        size_bytes: int,
        status: DocumentStatus,
    ) -> Document:
        now = datetime.now(UTC)
        self.document = Document(
            id=uuid4(),
            original_filename=original_filename,
            storage_key=storage_key,
            mime_type=mime_type,
            size_bytes=size_bytes,
            status=status,
            chunk_count=0,
            created_at=now,
            updated_at=now,
        )
        return self.document

    async def update_status(
        self,
        document: Document,
        status: DocumentStatus,
        *,
        error_message: str | None = None,
        page_count: int | None = None,
        chunk_count: int | None = None,
    ) -> Document:
        document.status = status
        document.error_message = error_message
        if page_count is not None:
            document.page_count = page_count
        if chunk_count is not None:
            document.chunk_count = chunk_count
        return document


class FakeEmbeddingProvider:
    async def embed_texts(
        self,
        texts: list[str],
        *,
        task: EmbeddingTask = "RETRIEVAL_DOCUMENT",
    ) -> list[list[float]]:
        assert task in {"RETRIEVAL_DOCUMENT", "RETRIEVAL_QUERY"}
        return [[0.1, 0.2, 0.3] for _ in texts]


class FakeLLMProvider:
    model = "fake-llm"

    async def stream_answer(self, messages: object) -> AsyncIterator[str]:
        del messages
        yield "Hello"
        yield " world"


class FakeChatRepository:
    def __init__(self, *, selected_documents: list[Document] | None = None) -> None:
        now = datetime.now(UTC)
        self.chat = Chat(id=uuid4(), title="Chat", created_at=now, updated_at=now)
        self.selected_documents = selected_documents or []
        self.messages: list[Message] = []

    async def list_all(self, *, limit: int = 100, offset: int = 0) -> tuple[list[Chat], int]:
        del limit, offset
        return [self.chat], 1

    async def get(self, chat_id: object) -> Chat | None:
        del chat_id
        return self.chat

    async def list_documents(self, chat_id: object) -> list[Document]:
        del chat_id
        return self.selected_documents

    async def list_ready_documents_by_ids(self, document_ids: list[object]) -> list[Document]:
        return [document for document in self.selected_documents if document.id in document_ids]

    async def replace_documents(self, chat: Chat, documents: list[Document]) -> list[Document]:
        del chat
        self.selected_documents = documents
        return documents

    async def create_message(
        self,
        *,
        chat_id: object,
        role: MessageRole,
        content: str,
        status: MessageStatus,
        model: str | None = None,
    ) -> Message:
        now = datetime.now(UTC)
        message = Message(
            id=uuid4(),
            chat_id=chat_id,
            role=role,
            content=content,
            status=status,
            model=model,
            created_at=now,
        )
        self.messages.append(message)
        return message

    async def update_message(
        self,
        message: Message,
        *,
        content: str | None = None,
        status: MessageStatus | None = None,
        latency_ms: int | None = None,
    ) -> Message:
        del latency_ms
        if content is not None:
            message.content = content
        if status is not None:
            message.status = status
        return message

    async def list_messages(self, chat_id: object) -> tuple[list[Message], int]:
        del chat_id
        return self.messages, len(self.messages)

    async def add_sources(
        self, *, message_id: object, chunk_scores: list[tuple[object, float]]
    ) -> None:
        del message_id, chunk_scores

    async def retrieve_chunks(
        self,
        *,
        document_ids: list[object],
        query_embedding: list[float],
        top_k: int,
        score_threshold: float | None,
    ) -> list[tuple[object, object, float]]:
        del document_ids, query_embedding, top_k, score_threshold
        return []


def test_chunking_produces_overlap() -> None:
    parsed = ParsedDocument(
        blocks=[
            ParsedBlock(text=f"Paragraph {index} " * 40, order_index=index)
            for index in range(6)
        ],
        page_count=None,
    )
    chunks = DocumentChunker(target_tokens=120, overlap_tokens=60).chunk(uuid4(), parsed)

    assert len(chunks) > 1
    assert chunks[0].metadata["block_end"] >= chunks[1].metadata["block_start"]


async def test_small_txt_creates_at_least_one_chunk() -> None:
    parsed = await TxtParser().parse(b"A very small document.", "small.txt")
    chunks = DocumentChunker(target_tokens=800, overlap_tokens=120).chunk(uuid4(), parsed)

    assert len(chunks) == 1
    assert chunks[0].token_count > 0


async def test_small_docx_creates_at_least_one_chunk() -> None:
    source = BytesIO()
    docx = DocxDocument()
    docx.add_paragraph("A very small DOCX document.")
    docx.save(source)

    parsed = await DocxParser().parse(source.getvalue(), "small.docx")
    chunks = DocumentChunker(target_tokens=800, overlap_tokens=120).chunk(uuid4(), parsed)

    assert len(chunks) == 1
    assert chunks[0].token_count > 0


async def test_text_pdf_creates_page_aware_chunk() -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "A small text PDF document.")
    data = pdf.tobytes()
    pdf.close()

    parsed = await PdfParser().parse(data, "small.pdf")
    chunks = DocumentChunker(target_tokens=800, overlap_tokens=120).chunk(uuid4(), parsed)

    assert len(chunks) == 1
    assert chunks[0].page_start == 1


async def test_upload_rejects_unsupported_file_type() -> None:
    service = DocumentService(FakeDocumentRepository(), Settings(), FakeStorage(), FakeJobQueue())

    with pytest.raises(UnsupportedMediaTypeError):
        await service.upload_document(
            original_filename="notes.exe",
            mime_type="application/octet-stream",
            data=b"content",
        )


async def test_upload_accepts_txt_and_queues_processing() -> None:
    storage = FakeStorage()
    queue = FakeJobQueue()
    service = DocumentService(FakeDocumentRepository(), Settings(), storage, queue)

    document = await service.upload_document(
        original_filename="notes.txt",
        mime_type="text/plain",
        data=b"hello",
    )

    assert document.status == DocumentStatus.PROCESSING
    assert storage.uploaded
    assert queue.enqueued == [("process_document", str(document.id))]


async def test_upload_enqueue_failure_marks_document_failed() -> None:
    repository = FakeDocumentRepository()
    service = DocumentService(repository, Settings(), FakeStorage(), FakeJobQueue(fail=True))

    with pytest.raises(DependencyUnavailableError):
        await service.upload_document(
            original_filename="notes.txt",
            mime_type="text/plain",
            data=b"hello",
        )

    assert repository.document is not None
    assert repository.document.status == DocumentStatus.FAILED
    assert repository.document.error_message == "Document processing job could not be queued"


async def test_chat_document_selection_requires_ready_documents() -> None:
    service = ChatService(
        FakeChatRepository(selected_documents=[]),
        Settings(),
        FakeEmbeddingProvider(),
        FakeLLMProvider(),
    )

    with pytest.raises(InvalidRequestError):
        await service.update_chat_documents(
            uuid4(),
            ChatDocumentSelectionUpdate(document_ids=[uuid4()]),
        )


async def test_message_stream_rejects_chat_with_no_selected_documents() -> None:
    service = ChatService(
        FakeChatRepository(selected_documents=[]),
        Settings(),
        FakeEmbeddingProvider(),
        FakeLLMProvider(),
    )

    with pytest.raises(InvalidRequestError):
        async for _ in service.stream_message(uuid4(), MessageCreate(content="Question?")):
            pass


async def test_message_stream_returns_sources_event_with_mocked_rag() -> None:
    now = datetime.now(UTC)
    document = Document(
        id=uuid4(),
        original_filename="notes.txt",
        storage_key="key",
        mime_type="text/plain",
        size_bytes=5,
        status=DocumentStatus.READY,
        created_at=now,
        updated_at=now,
    )
    service = ChatService(
        FakeChatRepository(selected_documents=[document]),
        Settings(),
        FakeEmbeddingProvider(),
        FakeLLMProvider(),
    )

    events = [
        event async for event in service.stream_message(uuid4(), MessageCreate(content="Question?"))
    ]

    assert events[0].startswith("event: sources")
    assert any("event: token" in event for event in events)
    assert events[-1].startswith("event: done")
