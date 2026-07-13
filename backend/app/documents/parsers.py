from dataclasses import dataclass, field
from pathlib import Path

import fitz
from docx import Document as DocxDocument


class DocumentParsingError(RuntimeError):
    """Raised when a document cannot be parsed into usable text."""


@dataclass(frozen=True, slots=True)
class ParsedBlock:
    text: str
    order_index: int
    page_number: int | None = None
    section_title: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    blocks: list[ParsedBlock]
    page_count: int | None
    metadata: dict[str, object] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks)


class DocumentParser:
    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        raise NotImplementedError


class PdfParser(DocumentParser):
    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        del filename
        try:
            pdf = fitz.open(stream=data, filetype="pdf")
        except Exception as exc:
            raise DocumentParsingError("PDF could not be opened") from exc

        blocks: list[ParsedBlock] = []
        try:
            for page_index in range(pdf.page_count):
                page = pdf.load_page(page_index)
                text = page.get_text("text").strip()
                if text:
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            order_index=len(blocks),
                            page_number=page_index + 1,
                        )
                    )
            if not blocks:
                raise DocumentParsingError(
                    "No extractable text found. Scanned PDFs require OCR, which is not enabled."
                )
            return ParsedDocument(
                blocks=blocks,
                page_count=pdf.page_count,
                metadata={"parser": "pymupdf"},
            )
        finally:
            pdf.close()


class TxtParser(DocumentParser):
    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        del filename
        text: str | None = None
        for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
            try:
                text = data.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if text is None or not text.strip():
            raise DocumentParsingError("TXT file does not contain extractable text")
        paragraphs = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n")]
        blocks = [
            ParsedBlock(text=paragraph, order_index=index)
            for index, paragraph in enumerate(paragraphs)
            if paragraph
        ]
        return ParsedDocument(blocks=blocks, page_count=None, metadata={"parser": "txt"})


class DocxParser(DocumentParser):
    async def parse(self, data: bytes, filename: str) -> ParsedDocument:
        del filename
        from io import BytesIO

        try:
            doc = DocxDocument(BytesIO(data))
        except Exception as exc:
            raise DocumentParsingError("DOCX file could not be opened") from exc

        blocks: list[ParsedBlock] = []
        current_section: str | None = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading"):
                current_section = text
            blocks.append(
                ParsedBlock(text=text, order_index=len(blocks), section_title=current_section)
            )

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(
                        ParsedBlock(
                            text=" | ".join(cells),
                            order_index=len(blocks),
                            section_title=current_section,
                        )
                    )

        if not blocks:
            raise DocumentParsingError("DOCX file does not contain extractable text")
        return ParsedDocument(blocks=blocks, page_count=None, metadata={"parser": "python-docx"})


def parser_for_document(filename: str, mime_type: str) -> DocumentParser:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or mime_type == "application/pdf":
        return PdfParser()
    if suffix == ".txt" or mime_type.startswith("text/"):
        return TxtParser()
    if (
        suffix == ".docx"
        or mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return DocxParser()
    raise DocumentParsingError("Unsupported document format")
